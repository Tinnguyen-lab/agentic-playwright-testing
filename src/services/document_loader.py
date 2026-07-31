"""Đọc tài liệu yêu cầu về text thuần cho agent. Hỗ trợ DOCX, PDF, TXT, MD."""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

SUPPORTED = (".docx", ".pdf", ".txt", ".md")


@dataclass
class LoadedDocument:
    text: str
    source_name: str
    doc_format: str
    n_units: int


def load_document(path: str | Path) -> LoadedDocument:
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".docx":
        return _load_docx(p)
    if ext == ".pdf":
        return _load_pdf(p)
    if ext in (".txt", ".md"):
        text = p.read_text(encoding="utf-8")
        return LoadedDocument(text, p.name, "text", text.count("\n") + 1)
    raise ValueError(f"Định dạng không hỗ trợ: {ext or '(không có đuôi)'}. Chỉ nhận {', '.join(SUPPORTED)}")


def _load_docx(p: Path) -> LoadedDocument:
    from docx import Document

    paras = [para.text for para in Document(str(p)).paragraphs]
    return LoadedDocument("\n".join(paras), p.name, "docx", len(paras))


def _load_pdf(p: Path) -> LoadedDocument:
    import fitz

    parts = []
    with fitz.open(str(p)) as doc:
        for i, page in enumerate(doc, start=1):
            parts.append(f"[trang {i}]\n{page.get_text()}")
    return LoadedDocument("\n".join(parts), p.name, "pdf", len(parts))
