"""Test parser tài liệu: DOCX/PDF/TXT/MD. Sinh file thật trong tmp_path, offline."""
import pytest

from src.services.document_loader import load_document


def test_load_txt(tmp_path):
    f = tmp_path / "r.md"
    f.write_text("Yêu cầu 1\nYêu cầu 2", encoding="utf-8")
    d = load_document(f)
    assert "Yêu cầu 1" in d.text
    assert d.doc_format == "text"
    assert d.source_name == "r.md"


def test_load_docx(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("Người dùng đăng nhập bằng email.")
    doc.add_paragraph("Hệ thống phản hồi nhanh.")
    f = tmp_path / "req.docx"
    doc.save(str(f))
    d = load_document(f)
    assert "đăng nhập" in d.text
    assert "phản hồi nhanh" in d.text
    assert d.doc_format == "docx"
    assert d.n_units == 2


def test_load_pdf(tmp_path):
    import fitz  # pymupdf
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Yeu cau kiem thu ABC")
    f = tmp_path / "req.pdf"
    doc.save(str(f))
    doc.close()
    d = load_document(f)
    assert "ABC" in d.text
    assert d.doc_format == "pdf"
    assert d.n_units == 1


def test_unsupported_ext(tmp_path):
    f = tmp_path / "x.csv"
    f.write_text("a,b", encoding="utf-8")
    with pytest.raises(ValueError):
        load_document(f)
